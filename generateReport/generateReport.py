#!/usr/bin/env python
# encoding: utf-8

import json
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from generatePic import generatePic

# 处理元数据
class Handle_metadata():

    def __init__(self,datafile):
        self.datafile = datafile
        self.metadata = self.get_metadata()

    def get_metadata(self):
        metadata = pd.read_json(self.datafile)
        return metadata

    # 获取按资产类型统计数据
    def get_assetTypeStatisData(self):
        metadata = self.metadata
        content = metadata[metadata.index.isin(["content"])]['res'].values[0]

#        print type(content["asset_analysis"]["asset_list"])
#        for k,v in content["asset_analysis"]["asset_list"][1].items():
#            print "========================="
#            print k
#            print v
        assetTypeStatis = content["asset_analysis"]["asset_statis"]["type"]
        assetTypeStatisResult = []
        for item in assetTypeStatis:
            temp = {}
            temp["type_name"] = item["type_name"]
            temp["asset_counts"] = item["asset_counts"]
            assetTypeStatisResult.append(temp)
        return assetTypeStatisResult

    # 获取cleandata
    def get_cleanData(self):
        cleandata = {}
        cleandata["assetTypeStatisData"] = self.get_assetTypeStatisData()
        return cleandata

class Generate_report_docx():
    def __init__(self, cleandata):
        self.cleandata = cleandata

    def generate(self):
        cleandata = self.cleandata
        document = Document()

        document.add_heading(u"工控脆弱性评估报告", 0)

        p = document.add_paragraph(u"生成工控脆弱性评估报告测试...")
#        p.style = 'ListBullet'
        # 概览
        document.add_heading(u"一. 概览", level=2)
        document.add_picture("img/overview.png", width=Inches(5.5))

        # 资产统计
        document.add_heading(u"二. 资产统计", level=2)
        document.add_heading(u"1.资产按产品类型统计", level=3)
        #document.add_paragraph(u"资产按产品类型统计", style="ListNumber")
        # 插入图表
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"

#        table.style = 'LightShading-Accent1'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u"产品类型"
        hdr_cells[1].text = u"数量"
        for item in cleandata["assetTypeStatisData"]:
            row_cells = table.add_row().cells
            row_cells[0].text = item["type_name"]
            row_cells[1].text = str(item["asset_counts"])

        document.add_heading(u"2. 资产类型分布图", level=3)
        document.add_picture("img/assetTypeStatis.png", width=Inches(6))


        # 资产脆弱性统计
        document.add_heading(u"三. 资产脆弱性统计", level=2)
        document.add_heading(u"1.脆弱性按厂商统计", level=3)
        #document.add_paragraph(u"资产按产品类型统计", style="ListNumber")
        # 插入图表
        table = document.add_table(rows=1, cols=2)
#        table.style = "Table Grid"
        table.style = "Medium List 1 Accent 3"
#        table.style = 'LightShading-Accent1'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u"产品类型"
        hdr_cells[1].text = u"数量"
        for item in cleandata["assetTypeStatisData"]:
            row_cells = table.add_row().cells
            row_cells[0].text = item["type_name"]
            row_cells[1].text = str(item["asset_counts"])

        document.add_heading(u"2. 脆弱性厂商分布图", level=3)
        document.add_picture("img/cuiruoxingChangshang.png", width=Inches(6))

        document.add_heading(u"3. 资产类型分布图", level=3)
        document.add_picture("img/hahaha.png", width=Inches(4.5))

        document.add_heading(u"4. 资产类型分布图", level=3)
        document.add_picture("img/haha.png", width=Inches(6))

        document.add_heading(u"4. 资产类型分布图", level=3)
        document.add_picture("img/piepie.png", width=Inches(5))

        document.add_heading(u"5. 资产类型分布图", level=3)
        document.add_picture("img/heihei.png", width=Inches(5))

        document.save("gongkongbaogao.docx")

if __name__ == "__main__":

    # 首先需要先生成所有需要的图片
#    gp = generatePic()
#    gp.generate_all_pic(all_pic_data)

    # 处理报告内容数据
    data_file = "data/report.json"
    dataObj = Handle_metadata(data_file)
    dataObj.get_assetTypeStatisData()

#    cleandata = dataObj.get_cleandata()

    # 生成docx格式报告
#    report = Generate_report_docx(cleandata)
#    report.generate()


