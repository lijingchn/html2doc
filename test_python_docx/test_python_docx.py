#!/usr/bin/env python
# encoding: utf-8

import json
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

# 处理元数据
class Handle_metadata():

    def __init__(self,datafile):
        self.datafile = datafile
        self.metadata = self.get_metadata()

    def get_metadata(self):
        metadata = pd.read_json(self.datafile)
        return metadata

    # 获取按资产类型统计数据
    def get_assetTypeStatisData(self):
        metadata = self.metadata
        content = metadata[metadata.index.isin(["content"])]['res'].values[0]
        for i in content:
            print i

        assetTypeStatis = content["asset_analysis"]["asset_statis"]["type"]
        assetTypeStatisResult = []
        for item in assetTypeStatis:
            temp = {}
            temp["type_name"] = item["type_name"]
            temp["asset_counts"] = item["asset_counts"]
            assetTypeStatisResult.append(temp)
        return assetTypeStatisResult

    # 获取cleandata
    def get_cleandata(self):
        cleandata = {}
        cleandata["assetTypeStatisData"] = self.get_assetTypeStatisData()
        return cleandata


#def get_meta_data(datafile):
#    meta_data = pd.read_json(datafile)
#    print meta_data
#    content = meta_data[meta_data.index.isin(["content"])]['res'].values[0]
#    print type(content)
#    for i in content:
#        print i
#    asset_type_statis = content["asset_analysis"]["asset_statis"]["type"]
#    print type(asset_type_statis)
#    print len(asset_type_statis)
#    asset_type_statis_result = []
#    for item in asset_type_statis:
#        temp = {}
#        temp["type_name"] = asset_type_statis[0]["type_name"]
#        temp["asset_counts"] = asset_type_statis[0]["asset_counts"]
#        asset_type_statis_result.append(temp)
#    return asset_type_statis_result

class Generate_report_docx():
    def __init__(self, cleandata):
        self.cleandata = cleandata

    def generate(self):
        cleandata = self.cleandata
        document = Document()

        document.add_heading(u"工控脆弱性评估报告", 0)

        p = document.add_paragraph(u"生成工控脆弱性评估报告测试...")
#        p.style = 'ListBullet'
        # 概览
        document.add_heading(u"一. 概览", level=2)
        document.add_picture("img/overview.png", width=Inches(5.5))

        # 资产统计
        document.add_heading(u"二. 资产统计", level=2)
        document.add_heading(u"1.资产按产品类型统计", level=3)
        #document.add_paragraph(u"资产按产品类型统计", style="ListNumber")
        # 插入图表
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"

#        table.style = 'LightShading-Accent1'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u"产品类型"
        hdr_cells[1].text = u"数量"
        for item in cleandata["assetTypeStatisData"]:
            row_cells = table.add_row().cells
            row_cells[0].text = item["type_name"]
            row_cells[1].text = str(item["asset_counts"])

        document.add_heading(u"2. 资产类型分布图", level=3)
        document.add_picture("img/assetTypeStatis.png", width=Inches(6))


        # 资产脆弱性统计
        document.add_heading(u"三. 资产脆弱性统计", level=2)
        document.add_heading(u"1.脆弱性按厂商统计", level=3)
        #document.add_paragraph(u"资产按产品类型统计", style="ListNumber")
        # 插入图表
        table = document.add_table(rows=1, cols=2)
#        table.style = "Table Grid"
        table.style = "Medium List 1 Accent 3"
#        table.style = 'LightShading-Accent1'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u"产品类型"
        hdr_cells[1].text = u"数量"
        for item in cleandata["assetTypeStatisData"]:
            row_cells = table.add_row().cells
            row_cells[0].text = item["type_name"]
            row_cells[1].text = str(item["asset_counts"])

        document.add_heading(u"2. 脆弱性厂商分布图", level=3)
        document.add_picture("img/cuiruoxingChangshang.png", width=Inches(6))

        document.add_heading(u"3. 资产类型分布图", level=3)
        document.add_picture("img/hahaha.png", width=Inches(4.5))

        document.add_heading(u"4. 资产类型分布图", level=3)
        document.add_picture("img/haha.png", width=Inches(6))

        document.add_heading(u"4. 资产类型分布图", level=3)
        document.add_picture("img/piepie.png", width=Inches(5))

        document.add_heading(u"5. 资产类型分布图", level=3)
        document.add_picture("img/heihei.png", width=Inches(5))

        document.save("gongkongbaogao.docx")

if __name__ == "__main__":
    data_file = "data/report.json"
    dataObj = Handle_metadata(data_file)
    cleandata = dataObj.get_cleandata()
    report = Generate_report_docx(cleandata)
    report.generate()


